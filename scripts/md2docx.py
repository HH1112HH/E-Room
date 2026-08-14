import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\AN NGOC\Desktop\CIT\E-Room\docs\roadmap.md"
OUT = r"C:\Users\AN NGOC\Desktop\CIT\E-Room\docs\Lo-trinh-phat-trien-MVP-E-Room.docx"

if len(sys.argv) > 1:
    OUT = sys.argv[1]

if len(sys.argv) > 2:
    SRC = sys.argv[2]

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

def set_heading_bg_color(paragraph, color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)

def add_text_with_formatting(paragraph, text):
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for t in tokens:
        if t.startswith("**") and t.endswith("**"):
            run = paragraph.add_run(t[2:-2])
            run.bold = True
        elif t.startswith("`") and t.endswith("`"):
            run = paragraph.add_run(t[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x25, 0x25)
        else:
            paragraph.add_run(t)

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)

def add_table(rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            text = cell.strip()
            if i == 0:
                run = table.cell(i, j).paragraphs[0].add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tcPr = table.cell(i, j)._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "2F5496")
                tcPr.append(shd)
            else:
                p = table.cell(i, j).paragraphs[0]
                add_text_with_formatting(p, text)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)

with open(SRC, encoding="utf-8") as f:
    raw = f.read()

lines = raw.split("\n")
i = 0
in_code = False
code_buf = []
in_table = False
table_buf = []

while i < len(lines):
    line = lines[i]

    if line.startswith("```"):
        if in_code:
            add_code_block(code_buf)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
        in_table = True
        table_buf = [line.strip("|").split("|")]
        i += 2
        continue
    if in_table and line.startswith("|"):
        table_buf.append(line.strip("|").split("|"))
        i += 1
        continue
    if in_table and not line.startswith("|"):
        add_table(table_buf)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        in_table = False
        table_buf = []

    if line.startswith("### "):
        h = doc.add_heading(level=3)
        add_text_with_formatting(h, line[4:])
        set_heading_bg_color(h, "D9E2F3")
        i += 1
        continue
    if line.startswith("## "):
        h = doc.add_heading(level=2)
        add_text_with_formatting(h, line[3:])
        set_heading_bg_color(h, "BDD7EE")
        i += 1
        continue
    if line.startswith("# "):
        h = doc.add_heading(level=1)
        add_text_with_formatting(h, line[2:])
        set_heading_bg_color(h, "8EAADB")
        i += 1
        continue

    if line.strip() == "---":
        i += 1
        continue

    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.italic = True
        add_text_with_formatting(p, line[2:])
        for run in p.runs:
            run.font.italic = True
        i += 1
        continue

    if re.match(r"^[-*] ", line):
        p = doc.add_paragraph(style="List Bullet")
        add_text_with_formatting(p, line[2:])
        i += 1
        continue

    if line.startswith("    "):
        p = doc.add_paragraph()
        run = p.add_run(line[4:])
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        i += 1
        continue

    if line.strip():
        p = doc.add_paragraph()
        add_text_with_formatting(p, line)
    i += 1

if in_table:
    add_table(table_buf)
if in_code:
    add_code_block(code_buf)

doc.save(OUT)
print("Saved:", OUT)